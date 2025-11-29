from fastapi import UploadFile
from docx import Document
import io

async def parse_docx(file: UploadFile) -> str:
    """
    Parse DOCX file and extract text content
    
    Args:
        file: UploadFile object containing DOCX
    
    Returns:
        Extracted text from DOCX
    """
    try:
        # Read file content
        content = await file.read()
        
        # Create a Document object
        doc_file = io.BytesIO(content)
        doc = Document(doc_file)
        
        # Extract text from all paragraphs
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Clean up the text
        text = text.strip()
        
        if not text:
            raise ValueError("No text could be extracted from the DOCX")
        
        return text
    
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")
    
    finally:
        # Reset file pointer for potential reuse
        await file.seek(0)